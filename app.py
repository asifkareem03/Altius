from flask import Flask, render_template, request, send_file
import os
import pandas as pd
from docx import Document as Documentx
from datetime import datetime, timedelta
import shutil
import tempfile
from docx2pdf import convert
from spire.doc import *
from spire.doc.common import *
from threading import Timer
import glob

app = Flask(__name__)

# Define the upload folder and allowed extensions   
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'xlsx'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def upload_file():
    return render_template('upload.html')

def process_excel(xls, output_file):
    # Read the Excel file
    # xls = pd.ExcelFile(input_file)
    
    # Initialize an empty DataFrame for the output
    output_df = pd.DataFrame(columns=['Name'])
    
    # Extract all dates from the sheet names
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name)
        date = pd.to_datetime(sheet_name, format='%d-%m-%Y').strftime('%d-%m-%Y')
        hour = df['Hour'][0]
        column_name = f'{date} ({hour})'
        output_df[column_name] = 'Absent'
        
        # Update 'Present' values for existing names
        output_df.loc[output_df['Name'].isin(df['Name']), column_name] = 'Present'

        # Add new names with 'yes' for the current date if they are not already in the DataFrame
        new_names = df[~df['Name'].isin(output_df['Name'])]['Name']
        if not new_names.empty:
            output_df = output_df._append(pd.DataFrame({'Name': new_names}), ignore_index=True)
            output_df.loc[output_df['Name'].isin(new_names), column_name] = 'Present'
    
    # Fill missing values with 'no'
    output_df.fillna('Absent', inplace=True)

    # Define a function to calculate the fee for each row
    def calculate_fee(row):
        fee = 0
        for col in row.index:
            if 'Present' in row[col] and col != 'Name':
                if '(2.0)' in col:
                    fee += 200 if row[col] == 'Present' else 0
                elif '(1.5)' in col:
                    fee += 150 if row[col] == 'Present' else 0
                elif '(1.0)' in col:
                    fee += 100 if row[col] == 'Present' else 0
                else:
                    fee += 250 if row[col] == 'Present' else 0
        return fee

    # Calculate fee based on the number of 'Present' occurrences in each row
    output_df['Fee'] = output_df.apply(calculate_fee, axis=1)

    # Save the output DataFrame to a new Excel file
    output_df.to_excel(output_file, index=False, sheet_name='Bill')

    return output_df


def fill_word_template(template_docx, data, output_folder):
    # Create a temporary directory to store the DOCX files
    temp_dir = tempfile.mkdtemp()

    # Load the template Word document
    doc = Documentx(template_docx)

    # Store references to frequently accessed elements
    tables = doc.tables
    invoice_number_cell = tables[0].cell(0, 1)
    current_date_cell = tables[0].cell(1, 1)
    name_cell = tables[0].cell(2, 1)
    coaching_fee_cell = tables[1].cell(1, 1)
    coaching_fee_amount_cell = tables[1].cell(1, 2)

    # Get previous month and year
    today = datetime.now()
    first_day_of_current_month = today.replace(day=1)
    last_day_of_previous_month = first_day_of_current_month - timedelta(days=1)
    previous_month = last_day_of_previous_month.strftime('%B')
    previous_month_number ='{:02d}'.format((last_day_of_previous_month.month - 1) % 12 + 1)
    previous_year = last_day_of_previous_month.year

    # Store original font color for final amount cell
    final_amount_cell = tables[2].cell(0, 0)
    original_font_color = final_amount_cell.paragraphs[0].runs[0].font.color.rgb

    for index, row in data.iterrows():
        name = row['Name']
        fee = row['Fee']
        discount = 0
        final_amount = fee - discount
        invoice_number = f"#{previous_month_number}{previous_year}{index+1}"

        # Fill in the data for the current invoice
        invoice_number_cell.text = invoice_number
        current_date_cell.text = datetime.now().strftime("%d-%m-%Y")
        name_cell.text = name

        coaching_fee_cell.text = f"Coaching Fee ({previous_month} {previous_year})"
        coaching_fee_amount_cell.text = str(fee)

        tables[1].cell(1, 4).text = str(final_amount)
        tables[1].cell(3, 2).text = f"Rs {fee}"
        tables[1].cell(3, 4).text = f"Rs {final_amount}"

        # Update the final amount cell with the calculated value
        final_amount_cell.text = f"Total Amount to be Paid: {final_amount}"
        final_amount_cell.paragraphs[0].runs[0].font.color.rgb = original_font_color

        # Save the filled Word document for the current name in the temporary directory
        temp_file_path = os.path.join(temp_dir, f'{name}_{str(previous_month)[:3]}_{str(previous_year)[-2:]}.docx')
        doc.save(temp_file_path)

        # Convert the DOCX file to PDF
        output_file_path = os.path.join(output_folder, f'{name}_{str(previous_month)[:3]}_{str(previous_year)[-2:]}.pdf')
        convert_to_pdf(temp_file_path, output_file_path)

    # Delete the temporary directory
    shutil.rmtree(temp_dir)

def convert_to_pdf(docx_file, pdf_file):
   document = Document()
   document.LoadFromFile(docx_file)
   document.SaveToFile(pdf_file, FileFormat.PDF)
   document.Close()

@app.route('/process', methods=['POST'])
def process_file():
    if 'file' not in request.files:
        return render_template('upload.html', message='No file part')
    
    file = request.files['file']
    
    if file.filename == '':
        return render_template('upload.html', message='No selected file')

    if file and allowed_file(file.filename):
        xls = pd.ExcelFile(file)

        if ((os.path.join(app.config['UPLOAD_FOLDER'], 'output'))):
            remove_files(((os.path.join(app.config['UPLOAD_FOLDER'], 'output'))))
        
        output_excel_file = os.path.join(app.config['UPLOAD_FOLDER'], 'output/output.xlsx')
        output_folder = os.path.join(app.config['UPLOAD_FOLDER'], 'output')

        # Process the uploaded Excel file
        data = process_excel(xls, output_excel_file)
        fill_word_template("name_month_year.docx", data, output_folder)
        
        zip_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'output.zip')
        shutil.make_archive(output_folder, 'zip', output_folder)
        return send_file(zip_file_path, as_attachment=True)
    else:
        return render_template('upload.html', message='Allowed file types are xlsx')
    
def remove_zip_file_delayed(file_path):
    try:
        os.remove(file_path)
    except Exception as e:
        print(f"Failed to delete {file_path}. Reason: {e}")

@app.after_request
def remove_zip_file(response):
    zip_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'output.zip')
    output_folder = os.path.join(app.config['UPLOAD_FOLDER'], 'output')
    if os.path.exists(zip_file_path):
        # Schedule the removal of the zip file after a short delay
        Timer(5, remove_zip_file_delayed, [zip_file_path]).start()
    if os.path.exists(output_folder):
        remove_files(output_folder)
        
    return response

def remove_files(dir):
    for filename in os.listdir(dir):
        file_path = os.path.join(dir, filename)
        try:
            if os.path.isfile(file_path):
                os.unlink(file_path)
        except Exception as e:
            print(f"Failed to delete {file_path}. Reason: {e}")

def remove_files_with_extension(directory, extension):
    files = glob.glob(os.path.join(directory, f'*.{extension}'))
    for file_path in files:
        try:
            os.remove(file_path)
        except Exception as e:
            print(f"Failed to delete {file_path}. Reason: {e}")

if __name__ == '__main__':
    app.run()
